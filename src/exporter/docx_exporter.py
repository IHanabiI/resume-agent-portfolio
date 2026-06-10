from __future__ import annotations

from io import BytesIO
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P


def markdown_to_docx_bytes(markdown_text: str) -> bytes:
    doc = Document()
    _apply_resume_base_styles(doc)
    _append_markdown_resume(doc, markdown_text)
    return _docx_to_bytes(doc)


def markdown_to_template_docx_bytes(markdown_text: str, template_bytes: bytes) -> bytes:
    """Fill an existing DOCX with resume text while keeping its shell.

    This keeps sections, styles, headers/footers, tables, and embedded media in
    place as much as python-docx allows. It is intentionally conservative: image
    only paragraphs are preserved and not used as text targets, while existing
    text paragraphs are replaced in document order.
    """

    doc = Document(BytesIO(template_bytes))
    if not markdown_text.strip():
        return _docx_to_bytes(doc)

    _apply_resume_base_styles(doc)
    _clear_template_body_keep_media(doc)
    _append_markdown_resume(doc, markdown_text)
    return _docx_to_bytes(doc)


def save_docx(markdown_text: str, output_dir: Path, filename: str = "tailored_resume.docx") -> Path:
    output_dir.mkdir(parents=True, exist_ok=True)
    path = output_dir / filename
    path.write_bytes(markdown_to_docx_bytes(markdown_text))
    return path


def _apply_resume_base_styles(doc: DocumentObject) -> None:
    section = doc.sections[0]
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)
    section.left_margin = Pt(45)
    section.right_margin = Pt(45)

    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(9.5)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(2)
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _append_markdown_resume(doc: DocumentObject, markdown_text: str) -> None:
    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("### "):
            _append_heading(doc, _clean_inline_markdown(line[4:].strip()), level=3)
        elif line.startswith("## "):
            _append_heading(doc, _clean_inline_markdown(line[3:].strip()), level=2)
        elif line.startswith("# "):
            _append_heading(doc, _clean_inline_markdown(line[2:].strip()), level=1)
        elif line.startswith("- "):
            _append_bullet(doc, line[2:].strip())
        elif line.startswith("* "):
            _append_bullet(doc, line[2:].strip())
        elif line.startswith("|"):
            _append_paragraph(doc, _clean_inline_markdown(line))
        else:
            _append_paragraph(doc, line)


def _append_heading(doc: DocumentObject, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6 if level <= 2 else 3)
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.color.rgb = RGBColor(31, 41, 55)
    run.font.size = Pt(14 if level == 1 else 11 if level == 2 else 10)


def _append_bullet(doc: DocumentObject, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Pt(12)
    paragraph.paragraph_format.first_line_indent = Pt(-6)
    paragraph.paragraph_format.space_after = Pt(1)
    _append_inline_runs(paragraph, text)


def _append_paragraph(doc: DocumentObject, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(1.5)
    _append_inline_runs(paragraph, text)


def _append_inline_runs(paragraph: Paragraph, text: str) -> None:
    text = text.replace("**", "__")
    parts = text.split("__")
    for index, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(_clean_inline_markdown(part))
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(9.5)
        run.bold = index % 2 == 1


def _clean_inline_markdown(text: str) -> str:
    return text.replace("`", "").replace("**", "").strip()


def _clear_template_body_keep_media(doc: DocumentObject) -> None:
    body = doc.element.body
    for child in list(body.iterchildren()):
        if child.tag.endswith("sectPr"):
            continue
        if _element_has_media(child):
            _clear_text_in_element(child)
            continue
        body.remove(child)


def _iter_paragraphs(parent: DocumentObject | _Cell) -> Iterable[Paragraph]:
    if isinstance(parent, DocumentObject):
        parent_elm = parent.element.body
    else:
        parent_elm = parent._tc

    for child in parent_elm.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            table = Table(child, parent)
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_paragraphs(cell)


def _element_has_media(element) -> bool:
    return bool(
        element.xpath(".//w:drawing")
        or element.xpath(".//w:pict")
        or element.xpath(".//a:blip")
    )


def _clear_text_in_element(element) -> None:
    for text_node in element.xpath(".//w:t"):
        text_node.text = ""


def _docx_to_bytes(doc: DocumentObject) -> bytes:
    stream = BytesIO()
    doc.save(stream)
    return stream.getvalue()
