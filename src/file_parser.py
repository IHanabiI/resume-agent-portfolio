from __future__ import annotations

from pathlib import Path

import pdfplumber
from docx import Document


def extract_text_from_upload(file_name: str, data: bytes) -> str:
    suffix = Path(file_name).suffix.lower()
    if suffix == ".txt":
        return data.decode("utf-8", errors="ignore")
    if suffix == ".pdf":
        return _extract_pdf(data)
    if suffix == ".docx":
        return _extract_docx(data)
    raise ValueError("仅支持 PDF、DOCX、TXT 文件。")


def _extract_pdf(data: bytes) -> str:
    from io import BytesIO

    text_parts: list[str] = []
    with pdfplumber.open(BytesIO(data)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(page_text)
    return "\n\n".join(text_parts)


def _extract_docx(data: bytes) -> str:
    from io import BytesIO

    doc = Document(BytesIO(data))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)

